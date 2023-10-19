
import json
import smtplib
import email
import smtplib
import datetime as dt
import icalendar
import pytz
import uuid
import ast
import json
import os
import re
import pandas as pd
import requests
from docx import Document
 
import ssl
from email import encoders
from email.message import EmailMessage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase

# Todos los correos han sido borrados, asi como las credenciales y API-Key por seguridad.
# Es un proyecto que solo interactua con PPM, outlook y chatgpt.

class BeFast():
    def __init__(self):
        self.correo_befast = ''
        with open('pwd.json') as f:
            dic = json.load(f)
            self.PSW = dic['pwd_befast']
            self.API_KEY = dic['api_key']
            self.URL_GPT = dic['url']

        f = open('peticion1.txt', "r")
        self.peticion1 = f.read().split('|| ')

        with open('PPM.json') as f:
            self.datosPPM = json.load(f)

        self.headers = {
            "Content-Type": "application/json",
            "api-key": self.API_KEY,
        }

        with open('convo1.json') as f:
            self.convo1 = json.load(f)

        with open('convo2.json') as f:
            self.convo2 = json.load(f)

        with open('convo3.json') as f:
            self.convo3 = json.load(f)

        with open('convo4.json') as f:
            self.convo4 = json.load(f)

        with open('convo5.json') as f:
            self.convo5 = json.load(f)

        with open('bodies.json') as f:
            self.bodies = json.load(f)

        self.smtpObj = None
        self.log_in_outlook()

    def log_in_outlook(self):
        try:
            self.smtpObj = smtplib.SMTP('smtp.office365.com', 587)
        except Exception as e:
            self.smtpObj = smtplib.SMTP_SSL('smtp.office365.com', 465)

        self.smtpObj.ehlo()
        self.smtpObj.starttls()
        self.smtpObj.login(self.correo_befast, self.PSW)
    
    def send_email(self, receptor, titulo, contenido, documento, firma):

        msg = EmailMessage()
        if documento == None:
            msg["Subject"] = titulo
            msg["From"] = self.correo_befast
            attendee_email = receptor[0]
            msg["To"] = attendee_email
            msg.set_content(contenido)
            self.smtpObj.send_message(msg = msg) # Or recipient@outlook
              
    def send_convo(self, receptores, titulo, contenido, fecha, hora, duracion):
        self.day = self.parse_date(fecha)
        start_hour = int(hora.split(':')[0])
        start_minute = int(hora.split(':')[1])
        # Timezone to use for our dates - change as needed
        tz = pytz.timezone("Europe/Madrid")
        start = tz.localize(dt.datetime.combine(self.day, dt.time(start_hour, start_minute, 0)))
        # Build the event itself
        cal = icalendar.Calendar()
        cal.add('prodid', '-//My calendar application//example.com//')
        cal.add('version', '2.0')
        cal.add('method', "REQUEST")
        event = icalendar.Event()
        for at in receptores:
            event.add('attendee', at)
        event.add('organizer', self.correo_befast)
        event.add('status', "confirmed")
        event.add('category', "Event")
        event.add('summary', titulo)
        event.add('description', contenido)
        event.add('location', "Barcelona")
        event.add('dtstart', start)
        event.add('dtend', tz.localize(dt.datetime.combine(self.day, dt.time(start_hour, start_minute + 30, 0))))
        event.add('dtstamp', tz.localize(dt.datetime.combine(dt.date.today(), dt.time(6, 0, 0))))
        event['uid'] = uuid.uuid4()#self.get_unique_id() # Generate some unique ID
        event.add('priority', 5)
        event.add('sequence', 1)
        event.add('created', tz.localize(dt.datetime.now()))

        # Add a reminder
        alarm = icalendar.Alarm()
        alarm.add("action", "DISPLAY")
        alarm.add('description', "Reminder")
        # The only way to convince Outlook to do it correctly
        #alarm.add("TRIGGER;RELATED=START", "-PT{0}H".format(reminder_hours))
        event.add_component(alarm)
        cal.add_component(event)

        # Build the email message and attach the event to it
        msg = MIMEMultipart("alternative")

        msg["Subject"] = titulo
        msg["From"] = self.correo_befast
        attendee_email = ', '.join(receptores)
        msg["To"] = attendee_email
        msg["Content-class"] = "urn:content-classes:calendarmessage"

        msg.attach(MIMEText(contenido))

        filename = "invite.ics"
        part = MIMEBase('text', "calendar", method="REQUEST", name=filename)
        part.set_payload( cal.to_ical() )
        email.encoders.encode_base64(part)
        part.add_header('Content-Description', filename)
        part.add_header("Content-class", "urn:content-classes:calendarmessage")
        part.add_header("Filename", filename)
        part.add_header("Path", filename)
        msg.attach(part)
        self.smtpObj.sendmail(msg["From"], [msg["To"]], msg.as_string())

    def parse_date(self, fecha):
        try:
            year = int(fecha.split('-')[2])
            month = int(fecha.split('-')[1])
            day = int(fecha.split('-')[0])
            ret = dt.date(year, month, day)
        except:
            print('El formato de la fecha del correo tiene que ser con el siguiente formato: DD-MM-AAAA')
        return ret
    
    def crearMensajeCorreo(self, mensajes, temperature):
        
        msg = []
        rol = 'system'
        for m in mensajes:
            if rol == 'system':
                msg.append({'role':'system', 'content' : m})
            else:
                msg.append({'role':'user', 'content' : m})

        data = {
            "messages": msg,
            "stream": False,
            # "max_tokens": 1000,
            "temperature": temperature,
        }
        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.dudasFuncionales = response.json()["choices"][0]["message"]["content"]
        demanda_ts = "TS"

        correo_ts = ( "Hola " + demanda_ts + ",\n" + "Has recibido la siguiente demanda. \n Titulo: " + self.datosPPM['demanda_titulo']  + "\n Descripción: " +
        self.datosPPM['demanda_descripcion'] + "\n  Objetivo: " + self.datosPPM['demanda_objetivo'] + "\n  Prioridad: " +
        self.datosPPM['demanda_prioridad'] + "\n  Peticionario: " + self.datosPPM['demanda_peticionario']  + "\n.  " +
        "Te propongo esta secuencia de dudas que deberíamos resolver: \n\n\n" + self.dudasFuncionales +
        "\n\n\n\n Quieres que cree una convocatoria para aterrizar los requerimientos funcionales con el PO? Contesta OK o contestame con el texto que quieres que ponga en la convocatoria.")
        
        return correo_ts
    
    def transcription(self):

        data = {
            "messages": [{"role": "system", "content": "Te encargas de generar conversaciones entre un product owner y un technical solutions sobre los requerimientos del proyecto.Los temas que se van a tratar son "+ self.dudasFuncionales}, {"role": "user", "content": ""}],
            "stream": False,
            "max_tokens": 1000,
            "temperature": 0.2,
        }

        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.transcripcionTSPO = response.json()["choices"][0]["message"]["content"]
    
    def resumen_transcripction(self):
        data = {
            "messages": [
                {"role": "system", "content": "eres un generador de resumenes. captura la informacion en bullets y no te dejes nada importante. Como dato de entrada te pasaré una transcripción de la llamada entre un technical solutions y un PO. Extrae la información relevante. "+ "Debes resumir la siguiente transcripción= "+ self.transcripcionTSPO
            }, {"role": "user", "content": "Resumen:"}],
            "stream": False,
            "max_tokens": 2000,
            "temperature": 0
        }

        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.resumenReunion = response.json()["choices"][0]["message"]["content"]

    def documentar_transcripcion(self) :
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run(self.transcripcionTSPO)
        day = str(dt.date.today().year) + str(dt.date.today().month) + str(dt.date.today().day) 
        doc.save("Transcription_meet_TS_PO_"+ day+".docx")
    
    def generar_docu_funcional(self):
        data = {

            "messages": [{"role": "system", "content": "Te encargas de crear un documento fucnional desde una conversacion transcrita entre  un product owner y un technical solutions sobre los requerimientos del proyecto. El documento debe tener las siquitentes secciones:"
            + "1. Introducción y objetivo \n"
            + "2. Requisitos funcionales \n"
            + "3. Requisitos no funcionales \n"
            "Para realizar el documento te paso la siguiente información:" + self.resumenReunion +
            "Descripción petición:" + self.datosPPM['demanda_descripcion'] +
            "Título petición:" + self.datosPPM['demanda_titulo']
            }, {"role": "user", "content": "Documento Funcional \n 1. Introducción y objetivo \n"
            }],
            "stream": False,
            # "max_tokens": 1000,
            "temperature": 0.2,
        }

        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.reqFuncional = response.json()["choices"][0]["message"]["content"]

    def generar_docu_tecnico(self):
        data = {
            "messages": [{"role": "system", "content":
            "Te encargas de crear un documento técnico sobre la elaboración de cuadros de mandos. El documento debe tener las siquitentes secciones:"
            + "1. Orignes de datos. \n"
            + "2. Transformaciones de datos necesarias  y modelo de datos en Qlik (dimensiones y métricas) \n"
            + "3. Seguridad de acceso a la información\n"
            + "4. Descripción del quadro de mando: Explicación de los diferentes paneles o secciones del quadro de mando, incluyendo la estructura, los indicadores y las métricas utilizadas en cada uno."
            + "Para realizar el documento te paso la siguiente información:" +
            "Documento Funcional"+ self.reqFuncional +
            "Descripción petición:" + self.datosPPM['demanda_descripcion'] +
            "Título petición:" + self.datosPPM['demanda_titulo'] +
            "Objetivo petición:" + self.datosPPM['demanda_objetivo']
            }, 
            {"role": "user", "content": "Documento Técnico \n 1. Orignes de datos \n"

        }],
        "stream": False,
        # "max_tokens": 1000,
        "temperature": 0.2,
        }
        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.reqTecnico = response.json()["choices"][0]["message"]["content"]
    
    
    def gen_doc(self):

        
        self.document = Document()
        self.document.add_heading(self.datosPPM['demanda_titulo'], 0)
        self.document.add_heading('0 - Información básica \n', level=1)
        p = self.document.add_paragraph('Descripción \n')
        records = (

            ('Prioridad Solicitante', self.datosPPM['demanda_prioridad']),

            ('Peticionario', self.datosPPM['demanda_peticionario']),

        )
        table = self.document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nombre proyecto'
        hdr_cells[1].text = self.datosPPM['demanda_titulo']

        for x, y in records:

            row_cells = table.add_row().cells
            row_cells[0].text = x
            row_cells[1].text = y

    

        self.document.add_heading('1 - Información funcional', level=1)
        p = self.document.add_paragraph(self.reqFuncional)
        self.document.add_page_break()

        self.document.add_heading('2 - Información técnica', level=1)
        p = self.document.add_paragraph(self.reqTecnico)
        self.document.add_page_break()

        #self.document.add_page_break()

        self.document.save("Documentacion_Funcional_Tecnico_v1.docx")

    def send_email_docu(self, body, tit, receptor,nombre_fichero):
        

        # Create a multipart message and set headers

        message = MIMEMultipart()

        message["From"] = self.correo_befast
        attendee_email = ', '.join(receptor)
        message["To"] = attendee_email
        message["Subject"] = tit
        
        # Add body to email

        message.attach(MIMEText(body, "plain"))
        filename = nombre_fichero  # In same directory as script

        # Open PDF file in binary mode

        with open(filename, "rb") as attachment:

            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())


        encoders.encode_base64(part)

        part.add_header(

            "Content-Disposition",
            f"attachment; filename= {filename}",

        )

        message.attach(part)
        text = message.as_string()
        context = ssl.create_default_context()

        self.smtpObj.sendmail(self.correo_befast, receptor, text) # Or recipient@outlook
    
    def plani_presupuesto(self):
        
        data = {
        "messages": [{"role": "system", "content":
        "Tienes que  realizar la planificación de proyecto en base al requerimiento técnico y funcional siguientes:" +
        "Documento Funcional"+ self.reqFuncional +
        "Documento Técnico"+ self.reqTecnico
        }, {"role": "user", "content": "A continuación te muestro la planificación estructurada por semanas:"
        }],
        "stream": False,
        "max_tokens": 1000,
        "temperature": 0.2,
        }

        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.document.add_heading('3 - Información económica y planificación', level=1)
        self.document.add_heading('Planificación', level=2)
        self.plani = response.json()["choices"][0]["message"]["content"]
        p = self.document.add_paragraph(self.plani)

        data = {

            "messages": [{"role": "system", "content":
            "Tienes que  rellenar la información económica y planificación de un proyecto de cuadro de mando. Debe ser una tabla de salida con los siguientes campos:"
            + "1. Valoración tentativa: Debe estar valorado en horas de trabajo. El campo debe ser un entero.\n"
            + "2. Disponibilidad de presupuesto: Debe ser un campo binario de TRUE/FALSE. \n"
            + "4. Proveedores a pedir valoración: Debes ser una lista con dos proveedores elegidos entre los siguientes [NTTData, SDG, MINDSAIT]   \n"
            "Para realizar el documento te paso la siguiente información:" +
            "Planificación"+ self.plani+
            "Documento Funcional"+ self.reqFuncional +
            "Documento Técnico"+ self.reqTecnico +
            "Fecha actual:" + "19/10/2023"

            }, {"role": "user", "content": ""

            }],
            "stream": False,
            # "max_tokens": 1000,
            "temperature": 0.2,
        }

        response = requests.post(self.URL_GPT, headers=self.headers, json=data)
        self.document.add_heading('Información económica', level=2)
        self.infEconomica = response.json()["choices"][0]["message"]["content"]
        p = self.document.add_paragraph(self.infEconomica)
        self.document.add_page_break()

        self.document.save("Documentacion_Funcional_Tecnico_v2.docx")
        # print(plani)
    
    def paso1y2(self):
        befast.send_email(receptor=[''], titulo='Recepcion PPM Demanda', contenido=correo, firma = 'BeFast', documento=None)
    
    def paso3(self):
        contenido = (self.convo1['saludo'] + '\n' + self.convo1['objetivo'] + " " +befast.datosPPM['demanda_titulo']+ '. \n\n ' + self.convo1['body'] + '\n\n' + befast.dudasFuncionales + '\n\n' + self.convo1['recordatorio'] + '\n\n' + self.convo1['despedida'] )
        befast.send_convo(receptores=self.convo1['convocados'], titulo=self.convo1['titulo'], contenido=contenido, fecha = self.convo1['fecha'], hora=self.convo1['hora'], duracion=self.convo1['duracion'])

    def paso4(self):

        befast.gen_doc()
        body = self.bodies['body_TS_1']
        tit = self.bodies['tit_TS_1']
        befast.send_email_docu(body,tit,['' ],"Documentacion_Funcional_Tecnico_v1.docx")

    def paso5(self):
        befast.gen_doc()
        befast.plani_presupuesto()
        body = self.bodies['body_TS_2']
        tit = self.bodies['tit_TS_2']
        befast.send_email_docu(body,tit,[''],"Documentacion_Funcional_Tecnico_v2.docx")
        

    def paso6(self):
        body = self.bodies['body_proveedores']
        tit = self.bodies['tit_proveedores']
        befast.send_email_docu(body,tit,[''],"Documentacion_Funcional_Tecnico_v1.docx")
        befast.send_email_docu(body,tit,[''],"Documentacion_Funcional_Tecnico_v1.docx")

    def paso7(self):
        contenido = (self.convo2['saludo'] + '\n' + self.convo2['objetivo'] + " " + befast.datosPPM['demanda_titulo']+ '. \n\n '  + self.convo2['despedida'] )
        befast.send_convo(receptores=self.convo2['convocados'], titulo=self.convo2['titulo'], contenido=contenido, fecha = self.convo2['fecha'], hora=self.convo2['hora'], duracion=self.convo2['duracion'])

        contenido = (self.convo3['saludo'] + '\n' + self.convo3['objetivo'] + " " +befast.datosPPM['demanda_titulo']+ '. \n\n '  + self.convo3['despedida'] )
        befast.send_convo(receptores=self.convo3['convocados'], titulo=self.convo3['titulo'], contenido=contenido, fecha = self.convo3['fecha'], hora=self.convo3['hora'], duracion=self.convo3['duracion'])

    def paso8(self):
        contenido = (self.convo4['saludo'] + '\n' + self.convo4['objetivo'] + " " + befast.datosPPM['demanda_titulo']+ '. \n\n '  + self.convo4['despedida'] )
        befast.send_convo(receptores=self.convo4['convocados'], titulo=self.convo4['titulo'], contenido=contenido, fecha = self.convo4['fecha'], hora=self.convo4['hora'], duracion=self.convo4['duracion'])

        contenido = (self.convo5['saludo'] + '\n' + self.convo5['objetivo'] + " " + befast.datosPPM['demanda_titulo']+ '. \n\n '  + self.convo5['despedida'] )
        befast.send_convo(receptores=self.convo5['convocados'], titulo=self.convo5['titulo'], contenido=contenido, fecha = self.convo5['fecha'], hora=self.convo5['hora'], duracion=self.convo5['duracion'])


    def paso9(self):
        body = self.bodies['correo_director']
        befast.send_email(receptor=[''], titulo='Aprobación PD ' + befast.datosPPM['demanda_titulo'], contenido=body, firma = 'BeFast', documento=None)


    def paso10(self):
        body = self.bodies['correo_rg']
        befast.send_email(receptor=[''], titulo='Aprobación PD ' + befast.datosPPM['demanda_titulo'], contenido=body, firma = 'BeFast', documento=None)


    def paso11(self):
        self.send_convo(receptores=befast.convo6['convocados'], titulo=befast.convo6['titulo'], contenido=befast.convo6['saludo'] + ' ' + befast.convo6['despedida'], fecha = befast.convo6['fecha'], hora=befast.convo6['hora'], duracion=befast.convo6['duracion'])


if __name__ == "__main__":

    befast = BeFast()
    correo = befast.crearMensajeCorreo(befast.peticion1, 0.2)
    befast.paso1y2()
    befast.paso3()


    
    befast.transcription()
    befast.documentar_transcripcion()
    befast.resumen_transcripction()
    befast.generar_docu_funcional()
    befast.generar_docu_tecnico()
    befast.paso4()

    befast.paso5()

    befast.paso6()
    befast.paso7()

    befast.paso8()
    befast.paso9()
    befast.paso10()
